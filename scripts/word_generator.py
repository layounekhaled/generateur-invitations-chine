#!/usr/bin/env python3
"""
Word Template Generator — Foshan Yingtabong Invitation
Modifies the uploaded .docx template directly by replacing placeholder data.
Then converts to PDF using LibreOffice.

Usage:
  python3 word_generator.py --data '{"lastName":"KRIBAA","firstName":"ABDELOUAHID",...}'
  python3 word_generator.py --data @request.json
"""

import sys
import os
import json
import shutil
import tempfile
import subprocess
import argparse
import re
from datetime import datetime, timedelta
from copy import deepcopy

# python-docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== Nationality Mapping ==============
NAT_MAP = {
    'Algeria': '阿尔及利亚', 'France': '法国', 'Morocco': '摩洛哥',
    'Tunisia': '突尼斯', 'Egypt': '埃及', 'Libya': '利比亚',
    'Mauritania': '毛里塔尼亚', 'Iraq': '伊拉克', 'Iran': '伊朗',
    'Turkey': '土耳其', 'Pakistan': '巴基斯坦', 'India': '印度',
    'Russia': '俄罗斯', 'Ukraine': '乌克兰', 'Nigeria': '尼日利亚',
    'Ghana': '加纳', 'Cameroon': '喀麦隆', 'Ethiopia': '埃塞俄比亚',
    'Kenya': '肯尼亚', 'South Africa': '南非', 'Indonesia': '印度尼西亚',
    'Malaysia': '马来西亚', 'Thailand': '泰国', 'Vietnam': '越南',
    'Philippines': '菲律宾', 'Bangladesh': '孟加拉国',
    'Senegal': '塞内加尔', 'Mali': '马里', 'Niger': '尼日尔',
    'Chad': '乍得', 'Sudan': '苏丹', 'Kazakhstan': '哈萨克斯坦',
    'Uzbekistan': '乌兹别克斯坦', 'Tanzania': '坦桑尼亚',
    'Congo': '刚果', 'Angola': '安哥拉', 'Mozambique': '莫桑比克',
    'Madagascar': '马达加斯加',
}

# ============== Embassy Mapping ==============
EMBASSY_MAP = {
    'Algeria': '中国驻阿尔及利亚使领馆签证中心',
    'France': '中国驻法国使领馆签证中心',
    'Morocco': '中国驻摩洛哥使领馆签证中心',
    'Tunisia': '中国驻突尼斯使领馆签证中心',
    'Egypt': '中国驻埃及使领馆签证中心',
    'Libya': '中国驻利比亚使领馆签证中心',
    'Mauritania': '中国驻毛里塔尼亚使领馆签证中心',
    'Iraq': '中国驻伊拉克使领馆签证中心',
    'Iran': '中国驻伊朗使领馆签证中心',
    'Turkey': '中国驻土耳其使领馆签证中心',
    'Pakistan': '中国驻巴基斯坦使领馆签证中心',
    'India': '中国驻印度使领馆签证中心',
    'Russia': '中国驻俄罗斯使领馆签证中心',
    'Ukraine': '中国驻乌克兰使领馆签证中心',
    'Nigeria': '中国驻尼日利亚使领馆签证中心',
    'Senegal': '中国驻塞内加尔使领馆签证中心',
    'Mali': '中国驻马里使领馆签证中心',
}


def date_to_cn(date_str):
    """Convert '2026-02-10' to '2026年2月10日'"""
    if not date_str:
        return ''
    try:
        dt = datetime.strptime(date_str, '%Y-%m-%d')
        return f'{dt.year}年{dt.month}月{dt.day}日'
    except:
        return date_str


def today_cn():
    """Return today's date in Chinese format"""
    n = datetime.now()
    return f'{n.year}年{n.month}月{n.day}日'


def date_short(date_str):
    """Convert '2026-02-10' to '2.10'"""
    if not date_str:
        return ''
    try:
        dt = datetime.strptime(date_str, '%Y-%m-%d')
        return f'{dt.month}.{dt.day}'
    except:
        return date_str


def generate_itinerary(arrival_str, departure_str):
    """Generate itinerary based on arrival and departure dates"""
    try:
        arrival = datetime.strptime(arrival_str, '%Y-%m-%d')
        departure = datetime.strptime(departure_str, '%Y-%m-%d')
    except:
        return []

    total_days = (departure - arrival).days + 1
    if total_days < 1:
        total_days = 1

    itinerary = []

    for i in range(total_days):
        current_date = arrival + timedelta(days=i)
        day_num = i + 1
        date_str = f'{current_date.month}.{current_date.day}'

        if i == 0:
            itinerary.append({
                'day': day_num,
                'date': date_str,
                'activity': '搭乘国际航班，飞往广州，夜宿航班。',
                'transport': '飞机',
                'meals': '无'
            })
        elif i == total_days - 1:
            itinerary.append({
                'day': day_num,
                'date': date_str,
                'activity': '酒店早餐后，送往广州白云国际机场，办理乘机手续。乘坐国际航班离开中国，返回阿尔及利亚，行程圆满结束。',
                'transport': '包车',
                'meals': '早'
            })
        elif i <= 2:
            itinerary.append({
                'day': day_num,
                'date': date_str,
                'activity': '考察广州中大纺织商圈，走访国际轻纺城和长江辅料城。',
                'transport': '包车',
                'meals': '早／午／晚'
            })
        elif i == 3:
            itinerary.append({
                'day': day_num,
                'date': date_str,
                'activity': '考察广州番禺服装产业带，走访沙溪商贸城和周边工厂。',
                'transport': '包车',
                'meals': '早／午／晚'
            })
        elif i == 4:
            itinerary.append({
                'day': day_num,
                'date': date_str,
                'activity': '东莞虎门服装批发市场采购。',
                'transport': '包车',
                'meals': '早／午／晚'
            })
        elif i >= total_days - 2:
            itinerary.append({
                'day': day_num,
                'date': date_str,
                'activity': '考察南沙保税仓。',
                'transport': '包车',
                'meals': '早／午／晚'
            })
        else:
            itinerary.append({
                'day': day_num,
                'date': date_str,
                'activity': '深度验厂与采购谈判。',
                'transport': '包车',
                'meals': '早／午／晚'
            })

    return itinerary


def set_cell_text_preserve_format(cell, text):
    """Replace text in a cell while preserving the first run's formatting."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        return

    for para in paragraphs:
        runs = list(para.runs)
        if not runs:
            continue

        # Get formatting from first run
        first_run = runs[0]
        rPr = first_run._element.find(qn('w:rPr'))

        # Remove all existing runs
        for run in runs:
            run._element.getparent().remove(run._element)

        # Remove proofErr elements
        for proof in para._element.findall(qn('w:proofErr')):
            para._element.remove(proof)

        # Create new run with preserved formatting
        new_run = OxmlElement('w:r')
        if rPr is not None:
            new_run.append(deepcopy(rPr))
        new_t = OxmlElement('w:t')
        new_t.text = text
        new_t.set(qn('xml:space'), 'preserve')
        new_run.append(new_t)
        para._element.append(new_run)


def set_date_cell(cell, date_str):
    """Replace a date cell with properly formatted Chinese date.
    Creates separate runs for numbers and Chinese characters with appropriate fonts."""
    if not date_str:
        return

    cn_date = date_to_cn(date_str)
    paragraphs = cell.paragraphs
    if not paragraphs:
        return

    for para in paragraphs:
        runs = list(para.runs)
        if not runs:
            continue

        # Get formatting from first run
        first_run = runs[0]
        rPr = first_run._element.find(qn('w:rPr'))

        # Remove all existing runs
        for run in runs:
            run._element.getparent().remove(run._element)

        # Remove proofErr elements
        for proof in para._element.findall(qn('w:proofErr')):
            para._element.remove(proof)

        # Parse the Chinese date and create runs with appropriate fonts
        # cn_date format: "2026年2月10日"
        parts = re.split(r'(年|月|日)', cn_date)
        for part in parts:
            if not part:
                continue
            new_run = OxmlElement('w:r')
            if rPr is not None:
                new_rPr = deepcopy(rPr)
                # For number parts, use Calibri; for Chinese chars, use SimSun
                if part in ('年', '月', '日'):
                    rFonts = new_rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        rFonts.set(qn('w:ascii'), 'SimSun')
                        rFonts.set(qn('w:eastAsia'), 'SimSun')
                        rFonts.set(qn('w:hAnsi'), 'SimSun')
                else:
                    rFonts = new_rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        rFonts.set(qn('w:ascii'), 'Calibri')
                        rFonts.set(qn('w:eastAsia'), 'Calibri')
                        rFonts.set(qn('w:hAnsi'), 'Calibri')
                new_run.append(new_rPr)
            new_t = OxmlElement('w:t')
            new_t.text = part
            new_t.set(qn('xml:space'), 'preserve')
            new_run.append(new_t)
            para._element.append(new_run)


def replace_text_in_paragraphs(doc, old_text, new_text):
    """Replace text across all paragraphs and tables in the document"""
    for para in doc.paragraphs:
        _replace_in_para(para, old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, old_text, new_text)


def _replace_in_para(para, old_text, new_text):
    """Replace text in a paragraph, handling runs that may be split"""
    full_text = para.text
    if old_text not in full_text:
        return

    # Simple case: text is in a single run
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return

    # Complex case: text spans multiple runs
    runs = list(para.runs)
    if not runs:
        return

    combined = ''
    run_map = []
    for run in runs:
        start = len(combined)
        combined += run.text
        end = len(combined)
        run_map.append((start, end, run))

    start_pos = combined.find(old_text)
    if start_pos == -1:
        return
    end_pos = start_pos + len(old_text)

    affected_runs = []
    for rs, re_, run in run_map:
        if rs < end_pos and re_ > start_pos:
            affected_runs.append((rs, re_, run))

    if not affected_runs:
        return

    first = True
    for rs, re_, run in affected_runs:
        if first:
            before = run.text[:max(0, start_pos - rs)]
            after_text = run.text[min(len(run.text), end_pos - rs):]
            run.text = before + new_text + after_text
            first = False
        else:
            if re_ > end_pos:
                after_text = run.text[min(len(run.text), end_pos - rs):]
                run.text = after_text
            else:
                run.text = ''


def _fix_split_dates_in_body(doc, new_date_cn):
    """Fix dates that are split across multiple runs in the body.
    For example: '2026' + '年' + '1' + '月' + '28' + '日'
    """
    for para in doc.paragraphs:
        full_text = para.text
        match = re.search(r'\d{4}年\d{1,2}月\d{1,2}日', full_text)
        if match and match.group() != new_date_cn:
            runs = list(para.runs)
            if not runs:
                continue

            # Get formatting from a non-empty run
            first_rPr = None
            for run in runs:
                rPr_elem = run._element.find(qn('w:rPr'))
                if rPr_elem is not None and run.text.strip():
                    first_rPr = deepcopy(rPr_elem)
                    break

            # Clear all runs
            for run in runs:
                run.text = ''

            # Set first non-empty run with new text
            if runs:
                runs[0].text = new_date_cn


def _create_itinerary_row(day_label, date_label, activity, transport, meals):
    """Create a new itinerary table row XML element"""
    new_tr = OxmlElement('w:tr')

    # Row properties
    trPr = OxmlElement('w:trPr')
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '480')
    trPr.append(trHeight)
    new_tr.append(trPr)

    # Column widths from the original table
    col_widths = [1199, 1318, 3669, 658, 1138]
    col_texts = [day_label, date_label, activity, transport, meals]

    for i, (width, text) in enumerate(zip(col_widths, col_texts)):
        tc = OxmlElement('w:tc')

        # Cell properties
        tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

        # Cell borders
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)
        tc.append(tcPr)

        # Cell paragraph
        p = OxmlElement('w:p')

        # Paragraph properties
        pPr = OxmlElement('w:pPr')
        wordWrap = OxmlElement('w:wordWrap')
        wordWrap.set(qn('w:val'), '0')
        pPr.append(wordWrap)
        autoSpaceDE = OxmlElement('w:autoSpaceDE')
        autoSpaceDE.set(qn('w:val'), '0')
        pPr.append(autoSpaceDE)
        autoSpaceDN = OxmlElement('w:autoSpaceDN')
        autoSpaceDN.set(qn('w:val'), '0')
        pPr.append(autoSpaceDN)

        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '20')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '239')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)

        jc = OxmlElement('w:jc')
        if i == 2:
            jc.set(qn('w:val'), 'both')
        else:
            jc.set(qn('w:val'), 'center')
        pPr.append(jc)

        pRPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '20')
        pRPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '18')
        pRPr.append(szCs)
        pPr.append(pRPr)
        p.append(pPr)

        # Run with text
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        rFonts = OxmlElement('w:rFonts')
        if i == 2:  # Activity - Chinese text
            rFonts.set(qn('w:ascii'), 'SimSun')
            rFonts.set(qn('w:eastAsia'), 'SimSun')
            rFonts.set(qn('w:hAnsi'), 'SimSun')
            rFonts.set(qn('w:hint'), 'eastAsia')
        elif i in [0, 1]:  # Day/Date
            rFonts.set(qn('w:ascii'), 'Calibri')
            rFonts.set(qn('w:eastAsia'), 'Calibri')
            rFonts.set(qn('w:hAnsi'), 'Calibri')
        else:
            rFonts.set(qn('w:ascii'), 'SimSun')
            rFonts.set(qn('w:eastAsia'), 'SimSun')
            rFonts.set(qn('w:hAnsi'), 'SimSun')
        rPr.append(rFonts)

        color = OxmlElement('w:color')
        color.set(qn('w:val'), '000000')
        rPr.append(color)

        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '20')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '18')
        rPr.append(szCs)

        r.append(rPr)

        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)

        p.append(r)
        tc.append(p)
        new_tr.append(tc)

    return new_tr


def modify_template(template_path, data, output_path):
    """Modify the Word template with the provided data"""
    doc = Document(template_path)

    # Extract data
    last_name = data.get('lastName', '')
    first_name = data.get('firstName', '')
    full_name = f'{last_name} {first_name}'.strip()
    sex = data.get('sex', 'M')
    nationality = data.get('nationality', 'Algeria')
    passport_number = data.get('passportNumber', '')
    arrival_date = data.get('arrivalDate', '')
    departure_date = data.get('departureDate', '')
    city_to_visit = data.get('cityToVisit', '广州、东莞等城市')
    inviter_company = data.get('inviterCompany', '佛山市盈达通外贸服务有限公司')

    # Derived values
    sex_cn = '男' if sex == 'M' else '女'
    nationality_cn = NAT_MAP.get(nationality, nationality)
    embassy = EMBASSY_MAP.get(nationality, '中国驻阿尔及利亚使领馆签证中心')
    arrival_cn = date_to_cn(arrival_date)
    departure_cn = date_to_cn(departure_date)
    today_str = today_cn()

    # ============== Page 1: Replace text in paragraphs ==============
    replace_text_in_paragraphs(doc, '中国驻阿尔及利亚使领馆签证中心', embassy)
    replace_text_in_paragraphs(doc, 'KRIBAA ABDELOUAHID', full_name)
    replace_text_in_paragraphs(doc, '176782160', passport_number)
    replace_text_in_paragraphs(doc, '广州、东莞等城市', city_to_visit)

    # Replace inviter company name
    replace_text_in_paragraphs(doc, '佛山市盈达通外贸服务有限公司', inviter_company)

    # ============== Page 1: Replace table cell values ==============
    tables = doc.tables

    if len(tables) >= 1:
        table = tables[0]
        rows = table.rows

        # Row indices (0-based):
        # Row 0: Header (被邀请人详细信息 | 具体内容)
        # Row 1: 姓名 | KRIBAA ABDELOUAHID
        # Row 2: 性别 | 男
        # Row 3: 国籍 | 阿尔及利亚
        # Row 4: 护照号码 | 176782160
        # Row 5: 预计入境日期 | 2026年2月10日
        # Row 6: 预计离境日期 | 2026年2月20日
        # Row 7: 访问城市 | 广州、东莞等城市

        if len(rows) >= 2:
            set_cell_text_preserve_format(rows[1].cells[1], full_name)

        if len(rows) >= 3:
            set_cell_text_preserve_format(rows[2].cells[1], sex_cn)

        if len(rows) >= 4:
            set_cell_text_preserve_format(rows[3].cells[1], nationality_cn)

        if len(rows) >= 5:
            set_cell_text_preserve_format(rows[4].cells[1], passport_number)

        if len(rows) >= 6:
            set_date_cell(rows[5].cells[1], arrival_date)

        if len(rows) >= 7:
            set_date_cell(rows[6].cells[1], departure_date)

        if len(rows) >= 8:
            set_cell_text_preserve_format(rows[7].cells[1], city_to_visit)

    # Fix split dates in body paragraphs (2026年1月28日 -> today)
    _fix_split_dates_in_body(doc, today_str)

    # ============== Page 2: Itinerary Table ==============
    if len(tables) >= 2:
        itinerary_table = tables[1]
        tbl_element = itinerary_table._tbl

        # Remove all existing rows
        existing_rows = tbl_element.findall(qn('w:tr'))

        # Keep the header row (first), remove the rest
        for tr in existing_rows[1:]:
            tbl_element.remove(tr)

        # Generate and add new itinerary rows
        itinerary = generate_itinerary(arrival_date, departure_date)

        for item in itinerary:
            new_row = _create_itinerary_row(
                day_label=f'Day{item["day"]}',
                date_label=item.get('date', ''),
                activity=item['activity'],
                transport=item['transport'],
                meals=item['meals']
            )
            tbl_element.append(new_row)

    # Save the modified document
    doc.save(output_path)
    return output_path


def convert_to_pdf(docx_path, pdf_path):
    """Convert .docx to .pdf using LibreOffice"""
    output_dir = os.path.dirname(pdf_path)
    try:
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            print(f'LibreOffice error: {result.stderr}', file=sys.stderr)
            return False

        # LibreOffice may save with a different name
        base_name = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        lo_output = os.path.join(output_dir, base_name)
        if lo_output != pdf_path and os.path.exists(lo_output):
            shutil.move(lo_output, pdf_path)

        return os.path.exists(pdf_path)
    except Exception as e:
        print(f'PDF conversion error: {e}', file=sys.stderr)
        return False


def main():
    parser = argparse.ArgumentParser(description='Generate invitation from Word template')
    parser.add_argument('--data', required=True, help='JSON data or @file.json')
    parser.add_argument('--output', default=None, help='Output path for .docx')
    parser.add_argument('--pdf', default=None, help='Also convert to PDF')
    parser.add_argument('--template', default=None, help='Path to .docx template')

    args = parser.parse_args()

    # Parse data
    data_str = args.data
    if data_str.startswith('@'):
        with open(data_str[1:], 'r') as f:
            data = json.load(f)
    else:
        data = json.loads(data_str)

    # Template path
    template_path = args.template or os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        'public', 'template.docx'
    )

    # Output path
    if args.output:
        output_path = args.output
    else:
        output_dir = tempfile.mkdtemp()
        safe_name = re.sub(r'[^a-zA-Z0-9]', '_', data.get('lastName', 'unknown') + '_' + data.get('firstName', 'invitee'))
        output_path = os.path.join(output_dir, f'invitation_{safe_name}.docx')

    # Generate
    modify_template(template_path, data, output_path)

    result = {'docx': output_path}

    # Convert to PDF if requested
    if args.pdf:
        pdf_path = args.pdf
        if convert_to_pdf(output_path, pdf_path):
            result['pdf'] = pdf_path

    # Output result as JSON
    print(json.dumps(result, ensure_ascii=False))


if __name__ == '__main__':
    main()
